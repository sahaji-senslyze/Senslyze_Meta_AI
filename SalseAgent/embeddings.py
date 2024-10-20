import os
import fitz  # PyMuPDF for PDFs
import re
import docx  # for DOCX files
from langchain_elasticsearch import ElasticsearchStore
from langchain_community.embeddings import HuggingFaceBgeEmbeddings
from uuid import uuid4
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from dotenv import load_dotenv


# Load environment variables from a specific file, such as .env.local
load_dotenv(dotenv_path=".env.local")


# jina_api_key = os.getenv("JINAAI_APIKEY")
ES_URL = "https://61fd8c4e92b648d4b058f085d5327a77.us-central1.gcp.cloud.es.io:443"
ES_API_KEY = "ZWlGS2Y1SUJpT25yNUgyQ3VhT1I6ZHMtTU5PTTVRWi0tRWNGaVpRQ1paZw=="

# Initialize embeddings and vector store
model_name = "BAAI/bge-small-en"
model_kwargs = {"device": "cpu"}
encode_kwargs = {"normalize_embeddings": True}
hf = HuggingFaceBgeEmbeddings(
    model_name=model_name, model_kwargs=model_kwargs, encode_kwargs=encode_kwargs
)
print(ES_API_KEY,ES_URL)

vector_store = ElasticsearchStore(
    es_url=ES_URL,
    index_name="booking_service",
    embedding=hf,
    es_api_key=ES_API_KEY
)
def extract_text_from_pdf(pdf_path):
    """Extracts text from a PDF file and includes page numbers."""
    with fitz.open(pdf_path) as pdf:
        content_with_page_numbers = []
        for page_num, page in enumerate(pdf, start=1):
            text = page.get_text("text")  # Extract text only
            clean_text = re.sub(r'[^A-Za-z0-9\s\n]+', '', text)  # Clean non-alphanumeric characters
            content_with_page_numbers.append((clean_text, page_num))
    return content_with_page_numbers

def extract_text_from_docx(docx_path):
    """Extracts text from a DOCX file (DOCX files have no concept of page numbers)."""
    doc = docx.Document(docx_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    clean_text = re.sub(r'[^A-Za-z0-9\s\n]+', '', text)  # Clean non-alphanumeric characters
    return [(clean_text, 1)]  # Return as a single chunk, no page numbers for DOCX

# Initialize the RecursiveCharacterTextSplitter
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=4000,  # Maximum size of each chunk
    chunk_overlap=100,  # Overlap between chunks
    length_function=len,  # Function to determine the length of the text
    is_separator_regex=False  # Whether to interpret the separator list as regex
)

def extract_metadata(file_path):
    """Extracts the file name and title from the document."""
    file_name = os.path.basename(file_path)
    title = os.path.splitext(file_name)[0]  # Use the file name (without extension) as the title.
    return file_name, title

def process_file(file_path):
    """Processes a file, extracts text and metadata, and performs chunking."""
    file_extension = file_path.split('.')[-1].lower()
    content_with_page_numbers = []

    if file_extension == 'pdf':
        content_with_page_numbers = extract_text_from_pdf(file_path)
    elif file_extension == 'docx':
        content_with_page_numbers = extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file format. Please use PDF or DOCX files.")

    # Extract metadata (file name and title)
    file_name, title = extract_metadata(file_path)

    # Chunk the extracted text
    documents = []
    for text, page_num in content_with_page_numbers:
        chunks = text_splitter.create_documents([text])  # Pass a list containing the text
        
        # Append metadata with page numbers to each chunk
        for chunk in chunks:
            document = Document(
                page_content=chunk.page_content,  # Use the chunked content
                metadata={
                    "file_name": file_name,
                    "file_path": file_path,
                    "page_number": page_num,  # Add page number metadata
                    "title": title  # Add title metadata
                }
            )
            documents.append(document)

    return documents

def process_directory(directory):
    """Processes all PDF and DOCX files in the specified directory."""
    documents_to_insert = []
    uuids = []

    # Traverse all files in the directory
    for file_name in os.listdir(directory):
        file_path = os.path.join(directory, file_name)

        # Process only PDF and DOCX files
        if file_name.endswith(".pdf") or file_name.endswith(".docx"):
            documents = process_file(file_path)  # Extract text and perform chunking
            for document in documents:
                documents_to_insert.append(document)
                uuids.append(str(uuid4()))  # Generate unique ID for each document

    # Add documents to the vector store
    vector_store.add_documents(documents=documents_to_insert, ids=uuids)

# Directory path to be processed (replace with your actual directory path)
directory_path = 'data/pdf'

# Process all files in the directory
process_directory(directory_path)